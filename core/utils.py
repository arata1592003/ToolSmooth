import os
import json
import docx
import re

import shutil

CRAWL_PATH = os.path.abspath("../output/crawl")
OUTPUT_PATH = os.path.abspath("../output")
DONE_PATH = os.path.abspath("../output_done")

def move_book_folders(site, book_id, to_done=True):
    """Di chuyển các folder liên quan đến truyện giữa output và output_done"""
    src_base = OUTPUT_PATH if to_done else DONE_PATH
    dst_base = DONE_PATH if to_done else OUTPUT_PATH
    
    # Các loại folder cần di chuyển
    sub_dirs = [
        os.path.join("translate", site, str(book_id)),
        os.path.join("smooth", site, str(book_id)),
        os.path.join("smooth", site, f"{book_id}_error"),
        os.path.join("smooth", site, f"{book_id}_skip"),
        os.path.join("crawl", site, str(book_id)),
    ]
    
    moved_any = False
    for sub in sub_dirs:
        src = os.path.join(src_base, sub)
        dst = os.path.join(dst_base, sub)
        
        if os.path.exists(src):
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            try:
                # Nếu folder đích đã tồn tại, xóa nó trước khi di chuyển để tránh lỗi
                if os.path.exists(dst):
                    shutil.rmtree(dst)
                shutil.move(src, dst)
                moved_any = True
            except Exception as e:
                print(f"❌ Lỗi khi di chuyển {src}: {e}")
                
    return moved_any

def read_docx(path):
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def write_smooth_docx(path, chapter_number, chapter_title, chapter_content):
    doc = docx.Document()

    # Tiêu đề chương
    title = f"Chương {chapter_number}: {chapter_title}"
    doc.add_heading(title, level=1)

    # Nội dung
    for line in chapter_content.split("\n"):
        doc.add_paragraph(line)

    doc.save(path)

def write_summary(path, chapter_number, summary_text):

    filename = os.path.join(path, f"Chuong_{chapter_number:05d}.txt")
    with open(filename, "w", encoding="utf-8") as f:
        f.write(summary_text)


def extract_number(filename):
    match = re.search(r'(\d{1,6})', filename)
    return int(match.group(1)) if match else None


def load_json(site, book_id):

    folder = os.path.join("../learn_output", site, book_id)
    os.makedirs(folder, exist_ok=True)

    path = os.path.join(folder, f"{book_id}.json")

    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {}

def get_last_translated(self):
    files = [
        f for f in os.listdir(self.output_folder)
        if f.lower().endswith(".docx") and re.search(r"\d+", f)
    ]
    if not files:
        return 0
    return max(int(re.search(r"\d+", f).group()) for f in files)

def load_rules():
    RULE_PATH = "./rules"
    rules = {}

    if not os.path.exists(RULE_PATH):
        os.makedirs(RULE_PATH)

    for filename in os.listdir(RULE_PATH):
        if filename.endswith(".txt"):
            rule_name = filename.replace(".txt", "")
            with open(os.path.join(RULE_PATH, filename), "r", encoding="utf-8") as f:
                rules[rule_name] = f.read()

    return rules

def load_sites():
    CONFIG_PATH = "../sites.json"
    if not os.path.exists(CONFIG_PATH):
        return {}

    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)
    
def load_category(site, book_id):
    book_info_path = os.path.join(CRAWL_PATH, f"{site}", f"{book_id}", "book_info.json")
    if os.path.exists(book_info_path):
        try:
            with open(book_info_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                category = data.get("category", "")
        except Exception:
            category = ""
    else:
        category = ""

    return category